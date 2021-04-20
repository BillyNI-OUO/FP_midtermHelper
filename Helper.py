import docx
import pandas as pd
import openpyxl
from tqdm import tqdm

########################################################

### EXCEL DATA EXTRACTION ###
# Excel file, input for the (number)
df = pd.read_excel(
    "./001-Attactment(2021-04-19).xlsx",
    usecols="B, D:FH",
    index_col="學號",
    engine='openpyxl'
)
# print(df)

# Replace data
# get the ID first
studentID = input("please input your student ID (ex: 107034058) : ")
dataOfId = df.loc[int(studentID)]
replace_data = {f"({i+1})": str(dataOfId.iloc[i]) for i in range(161)}

# print(replace_data)

########################################################

### WORD REPLACEMENT ###


def check_and_change(document, replace_dict):
    """
    遍历word中的所有 paragraphs，在每一段中发现含有key 的内容，就替换为 value 。
   （key 和 value 都是replace_dict中的键值对。）
    """
    # Paragraphs
    for para in tqdm(document.paragraphs):
        for key, value in replace_dict.items():
            if key in para.text:
                # print(key+"->"+value)
                para.text = para.text.replace(key, value)

    # Tables
    for table in tqdm(document.tables):
        for r in range(len(table.rows)):
            for c in range(len(table.columns)):
                if r == 0 or c == 0:
                    continue
                del_keyList = []
                for key, value in replace_dict.items():
                    if key in table.cell(r, c).text:
                        # print(key+"->"+value)
                        table.cell(r, c).text = table.cell(r, c).text.replace(
                            key, value)
                        del_keyList.append(key)

                for key in del_keyList:
                    replace_dict.pop(key, None)

    return document


# 创建文档对象,获得word文档
path = './000-Midterm Exam(2021-04-19).docx'
doc = docx.Document(path)

# 每一段的内容
for para in doc.paragraphs:
    print(para.text)

#  将想要替换的内容写成字典的形式，
#  dict = {"想要被替换的字符串": "新的字符串"}
doc = check_and_change(doc, replace_dict=replace_data)

# 每一段的编号、内容
# for i in range(len(doc.paragraphs)):
#    print(str(i), doc.paragraphs[i].text)

doc.save(f'{studentID}.docx')

########################################################
